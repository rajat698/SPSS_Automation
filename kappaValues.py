import docx
import os
import glob

folder_path = './kappaDocx'
files = glob.glob(os.path.join(folder_path, '*.docx'))

print(files )

new_doc = docx.Document()

num_rows = len(files) + 1
num_cols = 3
table = new_doc.add_table(rows=num_rows, cols=num_cols)

table.rows[0].cells[0].text = "Variable"
table.rows[0].cells[1].text = "kappa"
table.rows[0].cells[2].text = "p"

kappa = ''
p = ''

def read(doc):
    global kappa
    global p
    table = doc.tables[0]

    len_chi = len(table.rows[2].cells[1].text)

    kappa = table.rows[2].cells[2].text[:len_chi - 1]

    if kappa[0] == '.':
        kappa = "0" + kappa

    p = (table.rows[2].cells[5].text)

    print(f"kappa = {kappa}, p = {p}")

curr_row = 1

for i in files:
    doc = docx.Document(i)
    read(doc)

    table.rows[curr_row].cells[0].text = i[12:len(i) - 5]
    table.rows[curr_row].cells[1].text = kappa
    table.rows[curr_row].cells[2].text = p
    curr_row += 1

new_doc.save("kappaValues.docx")