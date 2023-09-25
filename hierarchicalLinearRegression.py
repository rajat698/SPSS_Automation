import docx
import os
import glob

def populate_coefficients_table(file):
    # Access the Coefficients table
    readTable_coeff = file.tables[7]

    # Determine the number of rows required
    variableCoeff = set()
    for row in readTable_coeff.rows:
        try:
            variableCoeff.add(int(row.cells[0].text))

        except ValueError:
            continue

    max_model = max(variableCoeff)
    max_starting_row = 0

    for i in range(len(readTable_coeff.rows)):
        if readTable_coeff.rows[i].cells[0].text == str(max_model):
            max_starting_row = i
            break

    write_row_list = readTable_coeff.rows[max_starting_row:len(readTable_coeff.rows)]

    num_rows = len(write_row_list) + 2
    num_cols = 7
    writeTable_coeff = write_doc.add_table(rows=num_rows, cols=num_cols)

    # Writing first default row
    row1 = writeTable_coeff.rows[0]
    cell1 = row1.cells[1]
    cell2 = row1.cells[2]
    cell1.merge(cell2)

    cell3 = row1.cells[3]
    cell4 = row1.cells[4]
    cell3.merge(cell4)

    cell5 = row1.cells[5]
    cell6 = row1.cells[6]
    cell5.merge(cell6)

    cell1.text = "Model 1a"
    cell3.text = "Model 1b"
    cell5.text = "Model 1c"

    # Writing second default row
    for i in range(1, 6, 2):
        writeTable_coeff.rows[1].cells[i].text = "ß"

    for i in range(2, 7 ,2):
        writeTable_coeff.rows[1].cells[i].text = "p"

    curr_row = 2

    #Writing variables
    for variable in write_row_list:
        writeTable_coeff.rows[curr_row].cells[0].text = variable.cells[1].text
        curr_row += 1

    #Writing "ß" and "p" values
    for row_write in writeTable_coeff.rows:
        for row_read in readTable_coeff.rows:
            if row_read.cells[1].text == row_write.cells[0].text:
                model = row_read.cells[0].text
                try:
                    column = 2 * int(model)
                
                except ValueError:
                    continue
                
                row_write.cells[column - 1].text = row_read.cells[4].text
                row_write.cells[column].text = row_read.cells[6].text


folder_path = './linearRegressionDocx'
files = glob.glob(os.path.join(folder_path, '*.docx'))

for i in files:
    read_doc = docx.Document(i)
    write_doc = docx.Document()
    populate_coefficients_table(read_doc)
    write_doc.save(f'{i[23:]}')