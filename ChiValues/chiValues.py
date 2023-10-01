import docx
import os
import glob

folder_path = './docx/chiDocx'
files = glob.glob(os.path.join(folder_path, '*.docx'))

new_doc = docx.Document()

num_rows = len(files) + 1
num_cols = 3
table = new_doc.add_table(rows=num_rows, cols=num_cols)

table.rows[0].cells[0].text = "Variable"
table.rows[0].cells[1].text = "X2"
table.rows[0].cells[2].text = "p"

chi = ''
p = ''

def read(doc):
    global chi
    global p
    table = doc.tables[0]

    len_chi = len(table.rows[2].cells[1].text)

    chi = table.rows[2].cells[1].text[:len_chi - 1]

    if chi[0] == '.':
        chi = "0" + chi

    p = (table.rows[2].cells[3].text)

    print(f"X2 = {chi}, p = {p}")


curr_row = 1

for i in files:
    doc = docx.Document(i)
    read(doc)

    table.rows[curr_row].cells[0].text = i[1:len(i) - 5]
    table.rows[curr_row].cells[1].text = chi
    table.rows[curr_row].cells[2].text = p
    curr_row += 1

new_doc.save("chiValues.docx")