import docx

file = './docx/ISTTestDocx/TOS T-test condition.docx'

read_file = docx.Document(file)
write_file = docx.Document()

groupStatistics = read_file.tables[1]

def PopulateGroupStatistics(read_doc, write_doc):

    

print(groupStatistics.rows[0].cells[0].text)
