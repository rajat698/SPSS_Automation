import docx
import toTwoDecimal

file = './docx/ISTTestDocx/TOS T-test 2.docx'


def initializeTable(read_doc, write_doc):

    num_groups = (((len(read_doc.tables[1].rows) - 2))//2)//3
    global num_rows
    num_rows = (len(read_doc.tables[1].rows) - 2)//2 + num_groups
    num_cols = 6

    global writeTable
    writeTable = write_doc.add_table(rows=num_rows, cols=num_cols)

    writeTable.rows[0].cells[
        1].text = "Behavioral Activation + Smoking Cessation Counseling M (SD)"
    writeTable.rows[0].cells[2].text = "Smoking Cessation Counseling M (SD)"
    writeTable.rows[0].cells[3].text = "t"
    writeTable.rows[0].cells[4].text = "p"
    writeTable.rows[0].cells[5].text = "d"

    row = 1
    heading_overLoad = 0
    global groupStatistics
    groupStatistics = read_doc.tables[1]
    i = 2

    while row < num_rows:
        if heading_overLoad == 3:
            writeTable.rows[row].cells[
                1].text = "Behavioral Activation + Smoking Cessation Counseling M (SD)"
            writeTable.rows[row].cells[2].text = "Smoking Cessation Counseling M (SD)"
            writeTable.rows[row].cells[3].text = "t"
            writeTable.rows[row].cells[4].text = "p"
            writeTable.rows[row].cells[5].text = "d"
            heading_overLoad = 0

        else:

            writeTable.rows[row].cells[0].text = groupStatistics.rows[i].cells[0].text
            i += 2

            heading_overLoad += 1

        row += 1


def populateTable():

    independentSamplesTable = read_doc.tables[2]
    independentSamplesEffectTable = read_doc.tables[3]

    print(independentSamplesEffectTable.rows[3].cells[3].text)

    readGSRow = 2
    readSTRow = 4
    readSETRow = 3
    writeRow = 1
    heading = 0

    while writeRow < num_rows and readGSRow < len(read_doc.tables[1].rows) and readSTRow < len(independentSamplesTable.rows):
        writeTable.rows[writeRow].cells[
            1].text = f"{toTwoDecimal.toTwoDecimal(groupStatistics.rows[readGSRow].cells[3].text)} ({toTwoDecimal.toTwoDecimal(groupStatistics.rows[readGSRow].cells[4].text)})"
        writeTable.rows[writeRow].cells[
            2].text = f"{toTwoDecimal.toTwoDecimal(groupStatistics.rows[readGSRow + 1].cells[3].text)} ({toTwoDecimal.toTwoDecimal(groupStatistics.rows[readGSRow + 1].cells[4].text)})"

        sig = independentSamplesTable.rows[readSTRow].cells[3].text

        if sig[0] == '<':
            sig = sig[1:]

        if float(sig) >= 0.05:
            writeTable.rows[writeRow].cells[3].text = independentSamplesTable.rows[readSTRow].cells[4].text
            writeTable.rows[writeRow].cells[4].text = independentSamplesTable.rows[readSTRow].cells[7].text

        else:
            writeTable.rows[writeRow].cells[3].text = independentSamplesTable.rows[readSTRow + 1].cells[4].text
            writeTable.rows[writeRow].cells[4].text = independentSamplesTable.rows[readSTRow + 1].cells[7].text
        
        d = toTwoDecimal.toTwoDecimal(independentSamplesEffectTable.rows[readSETRow].cells[3].text)
        
        if d[0] == '0':
            d = d[1:]
            if len(d) != 3:
                d = d + '0'
        
        elif d[0:2] == '-0':
            d = '-' + d[2:]
            if len(d) != 4:
                d = d + '0'

        writeTable.rows[writeRow].cells[5].text = independentSamplesEffectTable.rows[readSETRow].cells[3].text

        writeRow += 1
        heading += 1

        if heading == 3:
            writeRow += 1
            heading = 0

        readGSRow += 2
        readSTRow += 2
        readSETRow += 3


read_doc = docx.Document(file)
write_doc = docx.Document()

initializeTable(read_doc, write_doc)
populateTable()
write_doc.save('test2.docx')